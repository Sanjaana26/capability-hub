import streamlit as st
import boto3
import json
import os
import tempfile
from pathlib import Path
from dotenv import load_dotenv

load_dotenv()

bedrock = boto3.client(
    service_name="bedrock-runtime",
    region_name=os.getenv("AWS_DEFAULT_REGION", "us-east-1"),
    aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
    aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
    aws_session_token=os.getenv("AWS_SESSION_TOKEN"),
)

st.set_page_config(page_title="Capability Hub", page_icon="🧠", layout="wide")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

html, body, [class*="css"], .stApp {
    font-family: 'Inter', sans-serif !important;
    background-color: #f0ede6 !important;
}
.block-container { padding: 0 !important; max-width: 100% !important; }
#MainMenu, footer, header { visibility: hidden; }

.stSpinner > div { color: #0a2540 !important; }
[data-testid="stSpinner"] p { color: #0a2540 !important; }

[data-testid="stFileUploader"] section { background: white !important; }
[data-testid="stFileUploader"] span { color: #0a2540 !important; }
[data-testid="stFileUploader"] p { color: #0a2540 !important; }
[data-testid="stFileUploader"] small { color: #555 !important; }
[data-testid="stFileUploaderFileName"] { color: #0a2540 !important; }
[data-testid="stFileUploaderFileSize"] { color: #555 !important; }
[data-testid="stFileUploader"] button { background: #1a7a6e !important; color: white !important; border: none !important; }

.sp-topbar {
    background: #1a7a6e;
    color: white;
    padding: 0.6rem 2rem;
    font-size: 1rem;
    font-weight: 600;
    letter-spacing: 0.01em;
    display: flex;
    align-items: center;
    gap: 1rem;
}
.sp-topbar span { opacity: 0.85; font-size: 0.85rem; font-weight: 400; }

.orange-line {
    height: 5px;
    background: #c0392b;
    width: 60%;
    margin: 2rem 2.5rem 0 2.5rem;
}

.cap-body { padding: 0 2.5rem 2rem 2.5rem; max-width: 1200px; }

.cap-badge { color: #c0392b; font-size: 0.95rem; font-weight: 600; margin: 1.2rem 0 0.3rem 0; }
.cap-title { color: #0a2540; font-size: 2.4rem; font-weight: 700; margin: 0 0 0.5rem 0; line-height: 1.2; }
.cap-tagline { color: #444; font-size: 1rem; font-style: italic; margin: 0 0 1.5rem 0; }
.cap-divider { border: none; border-top: 1px solid #ccc; margin: 1rem 0 1.5rem 0; }

.sec-title { color: #c0392b; font-size: 1.3rem; font-weight: 700; margin: 2rem 0 0.8rem 0; padding-bottom: 0.4rem; border-bottom: 1px solid #ccc; }

.body-text { color: #222; font-size: 0.92rem; line-height: 1.8; margin-bottom: 0.8rem; }
.body-label { color: #222; font-size: 0.92rem; font-weight: 700; margin-bottom: 0.2rem; }

.dark-section {
    background: #1a7a6e;
    color: white;
    padding: 2rem 2.5rem;
    margin: 2rem 0;
}
.dark-section-title {
    color: #a8e6df;
    font-size: 0.8rem;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 0.08em;
    margin-bottom: 1.2rem;
    padding-bottom: 0.5rem;
    border-bottom: 1px solid rgba(255,255,255,0.25);
}
.dark-col-grid { display: grid; grid-template-columns: 1fr 1fr 1fr; gap: 2.5rem; }
.dark-col-body { color: #d4f0ec; font-size: 0.87rem; line-height: 1.8; }
.dark-col-body li { margin-bottom: 0.5rem; }
.dark-col-body strong { color: white; font-weight: 600; }

.pact-title { color: #c0392b; font-size: 1.8rem; font-weight: 700; text-align: center; margin: 2.5rem 0 1.5rem 0; }
.pact-grid { display: grid; grid-template-columns: 1fr 1fr 1fr 1fr; gap: 1.5rem; margin-bottom: 2rem; }
.pact-col { text-align: center; }
.pact-icon-wrap { font-size: 1.8rem; margin-bottom: 0.5rem; }
.pact-col-title { color: #c0392b; font-size: 1rem; font-weight: 700; margin-bottom: 0.8rem; }
.pact-col-body { color: #333; font-size: 0.85rem; line-height: 1.7; text-align: left; }
.pact-col-body li { margin-bottom: 0.4rem; }

.two-col { display: grid; grid-template-columns: 1fr 1fr; gap: 3rem; margin: 1rem 0 2rem 0; padding-left: 1rem; }
.obj-title { color: #c0392b; font-size: 1.4rem; font-weight: 700; margin: 2rem 0 1rem 0; text-align: center; }
.obj-item { margin-bottom: 1.2rem; padding-left: 1rem; border-left: 3px solid #c0392b; }
.obj-q { color: #222; font-size: 0.88rem; font-weight: 700; margin-bottom: 0.3rem; }
.obj-a { color: #444; font-size: 0.87rem; line-height: 1.7; }
.hook-item { color: #222; font-size: 0.88rem; line-height: 1.7; margin-bottom: 0.8rem; padding-left: 0.8rem; border-left: 3px solid #c0392b; }

.bottom-two-col { display: grid; grid-template-columns: 1fr 1fr; gap: 3rem; margin: 1rem 0 3rem 0; padding-left: 1rem; }
.faq-col-title { color: #c0392b; font-size: 1.4rem; font-weight: 700; text-align: center; margin-bottom: 1.2rem; }
.faq-item { margin-bottom: 1.2rem; padding-left: 0.8rem; border-left: 3px solid #1a7a6e; }
.faq-q { color: #222; font-size: 0.88rem; font-weight: 600; margin-bottom: 0.3rem; }
.faq-a { color: #444; font-size: 0.86rem; line-height: 1.7; }
.insight-item { color: #333; font-size: 0.87rem; line-height: 1.7; margin-bottom: 1rem; padding-left: 0.8rem; border-left: 3px solid #1a7a6e; }

.collaterals-section {
    border-top: 1px solid #ccc;
    padding: 2rem 0 1rem 0;
    text-align: center;
}
.collaterals-title { color: #c0392b; font-size: 1.6rem; font-weight: 700; margin-bottom: 1.5rem; }
.collaterals-actions { display: flex; gap: 1.5rem; justify-content: center; color: #555; font-size: 0.85rem; margin-bottom: 1.5rem; }
.collaterals-actions span { cursor: pointer; }
.comments-label { color: #222; font-size: 0.9rem; font-weight: 600; text-align: left; max-width: 600px; margin: 0 auto 0.8rem auto; }
.comment-box-wrap { display: flex; align-items: center; gap: 0.8rem; max-width: 600px; margin: 0 auto 2rem auto; }
.comment-avatar { width: 36px; height: 36px; border-radius: 50%; background: #ccc; display: flex; align-items: center; justify-content: center; font-size: 1.1rem; flex-shrink: 0; }
.comment-input-wrap { display: flex; flex: 1; border: 1px solid #ccc; border-radius: 4px; overflow: hidden; background: white; }
.comment-input { flex: 1; border: none; outline: none; padding: 0.5rem 0.8rem; font-size: 0.88rem; color: #333; background: white; }
.comment-post-btn { background: white; border: none; border-left: 1px solid #ccc; padding: 0.5rem 1rem; font-size: 0.88rem; color: #888; cursor: pointer; }
.comment-post-btn:hover { background: #f5f5f5; color: #333; }

.upload-wrap { padding: 3rem 2.5rem 1rem 2.5rem; max-width: 700px; margin: 0 auto; text-align: center; }
.upload-wrap h2 { color: #0a2540; font-size: 1.6rem; margin-bottom: 0.5rem; }
.upload-wrap p { color: #666; font-size: 0.95rem; }

.source-badge {
    display: inline-block;
    background: #1a7a6e;
    color: white;
    font-size: 0.75rem;
    font-weight: 600;
    padding: 0.2rem 0.8rem;
    border-radius: 20px;
    margin-bottom: 0.5rem;
}
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="sp-topbar">
    🧠 Capability Hub &nbsp;<span>| Agilisium</span>
</div>
""", unsafe_allow_html=True)

# ── S3 Functions ───────────────────────────────────────────────────────────────
def get_s3_client():
    return boto3.client(
        's3',
        region_name=os.getenv("AWS_DEFAULT_REGION", "us-east-1"),
        aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
        aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
        aws_session_token=os.getenv("AWS_SESSION_TOKEN"),
    )

def list_s3_documents():
    try:
        s3 = get_s3_client()
        response = s3.list_objects_v2(Bucket='vqb-doc-converter', Prefix='tech-docs/')
        files = []
        for obj in response.get('Contents', []):
            key = obj['Key']
            if key.endswith(('.pdf', '.docx', '.txt', '.md')):
                display_name = key.replace('tech-docs/', '')
                files.append({"key": key, "name": display_name})
        return files
    except Exception as e:
        st.error(f"❌ Could not connect to S3: {str(e)}")
        return []

def read_s3_document(key):
    try:
        s3 = get_s3_client()
        response = s3.get_object(Bucket='vqb-doc-converter', Key=key)
        file_bytes = response['Body'].read()
        suffix = Path(key).suffix.lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        text = ""
        try:
            if suffix == ".pdf":
                import PyPDF2
                with open(tmp_path, "rb") as f:
                    for page in PyPDF2.PdfReader(f).pages:
                        text += page.extract_text() or ""
            elif suffix in [".docx", ".doc"]:
                import docx
                text = "\n".join([p.text for p in docx.Document(tmp_path).paragraphs])
            else:
                with open(tmp_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
        finally:
            os.unlink(tmp_path)
        return text.strip()
    except Exception as e:
        st.error(f"❌ Could not read document from S3: {str(e)}")
        return ""

def extract_text(uploaded_file):
    suffix = Path(uploaded_file.name).suffix.lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.read())
        tmp_path = tmp.name
    text = ""
    try:
        if suffix == ".pdf":
            import PyPDF2
            with open(tmp_path, "rb") as f:
                for page in PyPDF2.PdfReader(f).pages:
                    text += page.extract_text() or ""
        elif suffix in [".docx", ".doc"]:
            import docx
            text = "\n".join([p.text for p in docx.Document(tmp_path).paragraphs])
        else:
            with open(tmp_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
    finally:
        os.unlink(tmp_path)
    return text.strip()


def call_claude(prompt):
    body = json.dumps({
        "anthropic_version": "bedrock-2023-05-31",
        "max_tokens": 8000,
        "messages": [{"role": "user", "content": prompt}]
    })
    response = bedrock.invoke_model(
        
        body=body
    )
    result = json.loads(response["body"].read())
    return result["content"][0]["text"]


def generate_capability_packet(doc_text):
    # ── CALL 1: Main content sections ──────────────────────────────────────────
    prompt1 = f"""You are a senior business consultant writing a formal Capability Packet for Agilisium's C-suite clients.

Read this technical document and generate detailed business content. Write for a CEO who has never seen this technology. Use plain English only — zero technical jargon.

YOU MUST write long, detailed paragraphs. Each paragraph field below needs AT LEAST 5 full sentences. Do not write short answers.

Return ONLY this JSON — no markdown, no backticks:

{{
  "title": "Professional capability name",
  "tagline": "One powerful sentence capturing the business transformation this delivers",
  "description": "Write exactly 5 sentences. Sentence 1: What this capability is in one line. Sentence 2: What specific business problem it solves. Sentence 3: Who uses it and how it fits into their daily work. Sentence 4: What makes this approach different from the status quo. Sentence 5: Why now is the right time to adopt this.",
  "current_pain": "Write exactly 5 sentences describing today's painful situation. Sentence 1: What does a typical day look like for people struggling with this problem? Sentence 2: What specific things go wrong and how often? Sentence 3: How much time and money is being wasted? Sentence 4: Who is frustrated and what opportunities are being missed? Sentence 5: What is the cumulative cost of continuing this way?",
  "business_impact": "Write exactly 5 sentences on consequences of NOT solving this. Sentence 1: The direct financial cost of inaction. Sentence 2: The competitive disadvantage this creates. Sentence 3: The operational risk this introduces. Sentence 4: The employee and customer experience impact. Sentence 5: What happens in 12 months if nothing changes?",
  "future_state": "Write exactly 5 sentences painting the transformed future. Sentence 1: What does a typical day look like after this solution is live? Sentence 2: What specific processes have changed and improved? Sentence 3: How have employees been freed up to do higher value work? Sentence 4: What new business opportunities are now possible? Sentence 5: What does the organization feel like from the inside?",
  "agilisium_role": "Write exactly 5 sentences on how Agilisium delivers this. Sentence 1: What is the overall delivery approach and philosophy? Sentence 2: What are the key phases of implementation? Sentence 3: What tools and methodologies does the team use? Sentence 4: What makes Agilisium's delivery low risk and reliable? Sentence 5: What does ongoing support and partnership look like?",
  "outcomes": [
    {{"title": "Outcome 1 — specific business result", "detail": "Write 4 sentences: what this outcome is, how it is achieved, what it means for daily operations, and the business value it creates."}},
    {{"title": "Outcome 2 — specific business result", "detail": "Write 4 sentences: what this outcome is, how it is achieved, what it means for daily operations, and the business value it creates."}},
    {{"title": "Outcome 3 — specific business result", "detail": "Write 4 sentences: what this outcome is, how it is achieved, what it means for daily operations, and the business value it creates."}}
  ],
  "primary_stakeholders": [
    "CEO / Executive Leadership — 2 sentences on exactly why they care and what they personally gain",
    "Operations Director — 2 sentences on exactly why they care and what they personally gain",
    "Finance Lead / CFO — 2 sentences on exactly why they care and what they personally gain",
    "IT / Technology Head — 2 sentences on exactly why they care and what they personally gain",
    "Business Analyst / Data Team Lead — 2 sentences on exactly why they care and what they personally gain"
  ],
  "secondary_stakeholders": [
    "End Users / Frontline Employees — 2 sentences on their interest and what improves for them",
    "Compliance / Risk Team — 2 sentences on their interest and what improves for them",
    "HR / People Operations — 2 sentences on their interest and what improves for them",
    "External Auditors / Regulators — 2 sentences on their interest",
    "Customers / End Clients — 2 sentences on how they benefit indirectly"
  ]
}}

TECHNICAL DOCUMENT:
{doc_text[:8000]}

Return ONLY the JSON. No markdown. No backticks."""

    # ── CALL 2: Supporting sections ─────────────────────────────────────────────
    prompt2 = f"""You are a senior business consultant writing the supporting sections of a Capability Packet for Agilisium.

Based on this technical document, write detailed business content for C-suite clients. Plain English only. No technical jargon.

Write long, detailed content. Each answer needs AT LEAST 3-4 full sentences.

Return ONLY this JSON — no markdown, no backticks:

{{
  "why_agilisium": [
    "Differentiator 1 title: Write 3 full sentences explaining this with specific proof points and examples of why this makes Agilisium the right partner.",
    "Differentiator 2 title: Write 3 full sentences explaining this with specific proof points and examples.",
    "Differentiator 3 title: Write 3 full sentences explaining this with specific proof points and examples.",
    "Differentiator 4 title: Write 3 full sentences explaining this with specific proof points and examples."
  ],
  "people": [
    "Role title — 2 sentences on what expertise this person brings and how it directly helps this engagement",
    "Role title — 2 sentences on expertise and value to the client",
    "Role title — 2 sentences on expertise and value to the client",
    "Role title — 2 sentences on expertise and value to the client",
    "Role title — 2 sentences on expertise and value to the client"
  ],
  "assets": [
    "Tool or framework name — 2 sentences on what it is and exactly how it accelerates delivery for the client",
    "Tool or framework name — 2 sentences on what it is and how it helps",
    "Tool or framework name — 2 sentences on what it is and how it helps",
    "Tool or framework name — 2 sentences on what it is and how it helps",
    "Tool or framework name — 2 sentences on what it is and how it helps"
  ],
  "case_studies": [
    "Case study 1: Write 4 sentences. Sentence 1: The client's situation and challenge. Sentence 2: What Agilisium specifically did. Sentence 3: The measurable result achieved. Sentence 4: Why this proves the approach works.",
    "Case study 2: Write 4 sentences describing a different situation and outcome with the same structure.",
    "Case study 3: Write 4 sentences describing another proof point."
  ],
  "thought_leadership": [
    "Insight title: Write 3 sentences on this strategic perspective and why it matters right now for business leaders.",
    "Insight title: Write 3 sentences on another strategic perspective that sets Agilisium apart.",
    "Insight title: Write 3 sentences on a third strategic perspective about the future of this space."
  ],
  "objections": [
    {{"q": "Is this solution too expensive for us right now?", "a": "Write 4 sentences: acknowledge the concern, reframe the cost as investment vs cost of inaction, provide a specific example of ROI, close with confidence."}},
    {{"q": "How long will implementation take and will it disrupt our operations?", "a": "Write 4 sentences: acknowledge the concern, explain the phased approach, describe how disruption is minimized, give a realistic timeline."}},
    {{"q": "We already have tools in place — why do we need this?", "a": "Write 4 sentences: acknowledge existing tools, explain the gap they leave, describe what this adds, explain why it complements rather than replaces."}},
    {{"q": "How do we know this will actually work for our specific situation?", "a": "Write 4 sentences: acknowledge the concern, reference similar client successes, explain the customization approach, offer a proof-of-concept step."}}
  ],
  "stakeholder_hooks": [
    "A powerful question about competitive disadvantage that makes a CEO want to act immediately — make it specific and urgent",
    "A powerful question about hidden costs or daily waste that would surprise a CFO — make it quantifiable",
    "A powerful question about employee productivity and the talent you are losing to frustration — make it personal",
    "A powerful question about what happens if a competitor solves this first — make it feel like a risk"
  ],
  "faqs": [
    {{"q": "How long does implementation typically take from start to go-live?", "a": "Write 4 sentences: give a realistic timeline, explain what happens in each phase, explain what affects the timeline, and close with confidence."}},
    {{"q": "How does this solution integrate with the systems we already have?", "a": "Write 4 sentences: explain the integration approach in plain language, describe what connects to what, explain how data flows, and reassure about compatibility."}},
    {{"q": "What does success look like in the first 90 days after go-live?", "a": "Write 4 sentences: describe specific early wins, explain what teams will notice changing, give concrete examples of improvements, and set realistic expectations."}},
    {{"q": "How do you ensure our data stays secure and compliant throughout?", "a": "Write 4 sentences: explain the security approach in plain language, describe compliance standards met, explain how data is protected at every step, and reassure about regulatory requirements."}},
    {{"q": "What ongoing support do we get after the solution goes live?", "a": "Write 4 sentences: describe the support model, explain what is included, describe how issues are handled, and explain how the solution evolves over time."}}
  ],
  "market_insights": [
    "Market insight 1: Write 4 sentences. Sentence 1: A key market trend making this capability urgent right now. Sentence 2: What competitors in the industry are already doing. Sentence 3: The cost of waiting or falling behind. Sentence 4: Why early adopters are pulling ahead.",
    "Market insight 2: Write 4 sentences on a regulatory or compliance shift that supports adopting this capability now and what it means for organizations that are not ready.",
    "Market insight 3: Write 4 sentences on where this space is heading in the next 2-3 years and what organizations who act now will be positioned to achieve."
  ]
}}

TECHNICAL DOCUMENT:
{doc_text[:8000]}

Return ONLY the JSON. No markdown. No backticks."""

    # Run both calls
    raw1 = call_claude(prompt1).strip()
    if raw1.startswith("```"):
        raw1 = raw1.split("```")[1]
        if raw1.startswith("json"):
            raw1 = raw1[4:]
    part1 = json.loads(raw1.strip())

    raw2 = call_claude(prompt2).strip()
    if raw2.startswith("```"):
        raw2 = raw2.split("```")[1]
        if raw2.startswith("json"):
            raw2 = raw2[4:]
    part2 = json.loads(raw2.strip())

    # Merge both results
    return {**part1, **part2}


# ── Session state ──────────────────────────────────────────────────────────────
if "packet" not in st.session_state:
    st.session_state.packet = None
if "doc_name" not in st.session_state:
    st.session_state.doc_name = ""
if "doc_source" not in st.session_state:
    st.session_state.doc_source = ""

# ── Upload / Select page ───────────────────────────────────────────────────────
if not st.session_state.packet:
    st.markdown("""
    <div class="orange-line"></div>
    <div class="upload-wrap">
        <h2>📄 Generate a Capability Packet</h2>
    </div>
    """, unsafe_allow_html=True)

    tab1, tab2 = st.tabs(["📦 Pick from S3 Bucket", "💻 Upload Manually"])

    with tab1:
        st.markdown("**Documents available in S3 bucket (`vqb-doc-converter`):**")
        s3_files = list_s3_documents()
        if s3_files:
            file_names = [f["name"] for f in s3_files]
            selected = st.selectbox("Select a document", file_names, label_visibility="collapsed")
            if st.button("🚀 Generate Capability Packet", type="primary", use_container_width=True):
                selected_key = next(f["key"] for f in s3_files if f["name"] == selected)
                with st.spinner("⏳ Reading from S3 and generating Capability Packet... (this may take 30-60 seconds)"):
                    try:
                        raw_text = read_s3_document(selected_key)
                        if raw_text:
                            packet = generate_capability_packet(raw_text)
                            st.session_state.packet = packet
                            st.session_state.doc_name = selected
                            st.session_state.doc_source = "S3"
                            st.rerun()
                    except Exception as e:
                        st.error(f"❌ Error: {str(e)}")
        else:
            st.info("No documents found in S3 bucket.")

    with tab2:
        uploaded_file = st.file_uploader("Choose file", type=["pdf", "docx", "txt", "md"], label_visibility="collapsed")
        if uploaded_file:
            with st.spinner("⏳ Generating Capability Packet... (this may take 30-60 seconds)"):
                try:
                    raw_text = extract_text(uploaded_file)
                    packet = generate_capability_packet(raw_text)
                    st.session_state.packet = packet
                    st.session_state.doc_name = uploaded_file.name
                    st.session_state.doc_source = "Manual"
                    st.rerun()
                except Exception as e:
                    st.error(f"❌ Error: {str(e)}")

# ── Render Capability Packet ───────────────────────────────────────────────────
else:
    p = st.session_state.packet

    st.markdown('<div class="orange-line"></div>', unsafe_allow_html=True)

    source = st.session_state.doc_source
    source_label = "📦 Source: S3 Bucket" if source == "S3" else "💻 Source: Manual Upload"
    st.markdown(f'<div style="padding: 0.5rem 2.5rem;"><span class="source-badge">{source_label} — {st.session_state.doc_name}</span></div>', unsafe_allow_html=True)

    st.markdown(f"""
    <div class="cap-body">
        <div class="cap-badge">Capability Packet</div>
        <div class="cap-title">{p.get('title', '')}</div>
        <div class="cap-tagline">{p.get('tagline', '')}</div>
        <hr class="cap-divider">
        <div class="sec-title">Description of the capability</div>
        <div class="body-text">{p.get('description', '')}</div>
        <div class="body-label">Current state pain:</div>
        <div class="body-text">{p.get('current_pain', '')}</div>
        <div class="body-label">Impact on the business:</div>
        <div class="body-text">{p.get('business_impact', '')}</div>
        <div class="body-label">Future state:</div>
        <div class="body-text">{p.get('future_state', '')}</div>
        <div class="body-label">Agilisium role:</div>
        <div class="body-text">{p.get('agilisium_role', '')}</div>
    </div>
    """, unsafe_allow_html=True)

    outcomes = p.get("outcomes", [])
    primary = p.get("primary_stakeholders", [])
    secondary = p.get("secondary_stakeholders", [])
    why = p.get("why_agilisium", [])

    outcomes_html = ""
    for o in outcomes:
        outcomes_html += f"<div style='margin-bottom:1.2rem'><strong>{o.get('title','')}</strong><br><span style='color:#d4f0ec;font-size:0.85rem;line-height:1.7'>{o.get('detail','')}</span></div>"

    primary_html = "".join([f"<li>{s}</li>" for s in primary])
    secondary_html = "".join([f"<li>{s}</li>" for s in secondary])
    why_html = "".join([f"<div style='margin-bottom:1rem;color:#d4f0ec;font-size:0.87rem;line-height:1.7'>{w}</div>" for w in why])

    st.markdown(f"""
    <div class="dark-section">
        <div class="dark-col-grid">
            <div>
                <div class="dark-section-title">Business value and success measures</div>
                <div class="dark-col-body">{outcomes_html}</div>
            </div>
            <div>
                <div class="dark-section-title">Primary Stakeholders</div>
                <div class="dark-col-body"><ul>{primary_html}</ul></div>
                <div class="dark-section-title" style="margin-top:1.5rem">Secondary Stakeholders</div>
                <div class="dark-col-body"><ul>{secondary_html}</ul></div>
            </div>
            <div>
                <div class="dark-section-title">Why Agilisium</div>
                {why_html}
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    people = p.get("people", [])
    assets = p.get("assets", [])
    cases = p.get("case_studies", [])
    tl = p.get("thought_leadership", [])

    people_html = "".join([f"<li>{x}</li>" for x in people])
    assets_html = "".join([f"<li>{x}</li>" for x in assets])
    cases_html = "".join([f"<li>{x}</li>" for x in cases])
    tl_html = "".join([f"<li>{x}</li>" for x in tl])

    st.markdown(f"""
    <div class="cap-body">
        <div class="pact-title">What Agilisium can offer (PACT)</div>
        <div class="pact-grid">
            <div class="pact-col">
                <div class="pact-icon-wrap">👥</div>
                <div class="pact-col-title">People</div>
                <div class="pact-col-body"><ul>{people_html}</ul></div>
            </div>
            <div class="pact-col">
                <div class="pact-icon-wrap">⚙️</div>
                <div class="pact-col-title">Assets</div>
                <div class="pact-col-body"><ul>{assets_html}</ul></div>
            </div>
            <div class="pact-col">
                <div class="pact-icon-wrap">📋</div>
                <div class="pact-col-title">Case Studies</div>
                <div class="pact-col-body"><ul>{cases_html}</ul></div>
            </div>
            <div class="pact-col">
                <div class="pact-icon-wrap">💡</div>
                <div class="pact-col-title">Thought Leadership</div>
                <div class="pact-col-body"><ul>{tl_html}</ul></div>
            </div>
        </div>
    """, unsafe_allow_html=True)

    objections = p.get("objections", [])
    hooks = p.get("stakeholder_hooks", [])

    obj_html = ""
    for obj in objections:
        obj_html += f"""
        <div class="obj-item">
            <div class="obj-q">"{obj.get('q','')}"</div>
            <div class="obj-a">{obj.get('a','')}</div>
        </div>"""

    hooks_html = ""
    for h in hooks:
        hooks_html += f'<div class="hook-item">{h}</div>'

    st.markdown(f"""
        <div class="two-col">
            <div>
                <div class="obj-title">Objection Handling</div>
                {obj_html}
            </div>
            <div>
                <div class="obj-title">Stakeholder Hooks</div>
                {hooks_html}
            </div>
        </div>
    """, unsafe_allow_html=True)

    faqs = p.get("faqs", [])
    insights = p.get("market_insights", [])

    faqs_html = ""
    for faq in faqs:
        faqs_html += f"""
        <div class="faq-item">
            <div class="faq-q">{faq.get('q','')}</div>
            <div class="faq-a">{faq.get('a','')}</div>
        </div>"""

    insights_html = ""
    for ins in insights:
        insights_html += f'<div class="insight-item">{ins}</div>'

    st.markdown(f"""
        <div class="bottom-two-col">
            <div>
                <div class="faq-col-title">FAQs</div>
                {faqs_html}
            </div>
            <div>
                <div class="faq-col-title">Key Market Insights</div>
                {insights_html}
            </div>
        </div>
    """, unsafe_allow_html=True)

    st.markdown("""
        <div class="collaterals-section">
            <div class="collaterals-title">Collaterals</div>
            <hr style="border:none;border-top:1px solid #ccc;margin:0 auto 1.2rem auto;max-width:600px">
            <div class="collaterals-actions">
                <span>👍 Like</span>
                <span>💬 Comment</span>
                <span>👁 Views</span>
                <span>🔖 Save for later</span>
            </div>
            <div class="comments-label">Comments</div>
            <div class="comment-box-wrap">
                <div class="comment-avatar">👤</div>
                <div class="comment-input-wrap">
                    <input class="comment-input" type="text" placeholder="Add a comment. Type @ to mention someone" />
                    <button class="comment-post-btn">Post</button>
                </div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    download_html = f"""
    <html><head><meta charset='utf-8'>
    <style>
        body {{ font-family: Inter, sans-serif; padding: 2rem 3rem; background: #f0ede6; color: #222; }}
        h1 {{ color: #0a2540; font-size: 2rem; }}
        h2 {{ color: #c0392b; font-size: 1.2rem; border-bottom: 1px solid #ccc; padding-bottom: 0.3rem; margin-top: 2rem; }}
        .label {{ font-weight: 700; margin-top: 1rem; }}
        .body {{ margin-bottom: 0.8rem; line-height: 1.8; }}
        .badge {{ background: #0a2540; color: white; padding: 0.2rem 0.8rem; border-radius: 20px; font-size: 0.8rem; }}
        ul {{ padding-left: 1.2rem; line-height: 1.8; }}
        .obj {{ border-left: 3px solid #c0392b; padding-left: 1rem; margin-bottom: 1rem; }}
        .faq {{ border-left: 3px solid #1a7a6e; padding-left: 1rem; margin-bottom: 1rem; }}
        .hook {{ border-left: 3px solid #c0392b; padding-left: 1rem; margin-bottom: 0.8rem; }}
        .insight {{ border-left: 3px solid #1a7a6e; padding-left: 1rem; margin-bottom: 0.8rem; }}
        .grid3 {{ display: grid; grid-template-columns: 1fr 1fr 1fr; gap: 2rem; background: #1a7a6e; color: white; padding: 2rem; margin: 1rem 0; }}
        .grid3 h3 {{ color: #a8e6df; font-size: 0.8rem; text-transform: uppercase; letter-spacing: 0.08em; border-bottom: 1px solid rgba(255,255,255,0.2); padding-bottom: 0.4rem; }}
        .grid3 p, .grid3 li {{ color: #d4f0ec; font-size: 0.87rem; line-height: 1.8; }}
        .pact {{ display: grid; grid-template-columns: 1fr 1fr 1fr 1fr; gap: 1.5rem; margin: 1rem 0; }}
        .pact-col h3 {{ color: #c0392b; font-size: 0.95rem; }}
        .two {{ display: grid; grid-template-columns: 1fr 1fr; gap: 2rem; }}
    </style></head><body>
    <span class='badge'>Capability Packet</span>
    <h1>{p.get('title','')}</h1>
    <p><em>{p.get('tagline','')}</em></p>
    <hr>
    <h2>Description of the capability</h2>
    <p class='body'>{p.get('description','')}</p>
    <p class='label'>Current state pain:</p><p class='body'>{p.get('current_pain','')}</p>
    <p class='label'>Impact on the business:</p><p class='body'>{p.get('business_impact','')}</p>
    <p class='label'>Future state:</p><p class='body'>{p.get('future_state','')}</p>
    <p class='label'>Agilisium role:</p><p class='body'>{p.get('agilisium_role','')}</p>
    <div class='grid3'>
        <div>
            <h3>Business Value and Success Measures</h3>
            {''.join([f"<p><strong style='color:white'>{o.get('title','')}</strong><br>{o.get('detail','')}</p>" for o in p.get('outcomes',[])])}
        </div>
        <div>
            <h3>Primary Stakeholders</h3>
            <ul>{''.join([f'<li>{s}</li>' for s in p.get('primary_stakeholders',[])])}</ul>
            <h3 style='margin-top:1rem'>Secondary Stakeholders</h3>
            <ul>{''.join([f'<li>{s}</li>' for s in p.get('secondary_stakeholders',[])])}</ul>
        </div>
        <div>
            <h3>Why Agilisium</h3>
            {''.join([f'<p>{w}</p>' for w in p.get('why_agilisium',[])])}
        </div>
    </div>
    <h2 style='text-align:center;border:none'>What Agilisium can offer (PACT)</h2>
    <div class='pact'>
        <div class='pact-col'><h3>👥 People</h3><ul>{''.join([f'<li>{x}</li>' for x in p.get('people',[])])}</ul></div>
        <div class='pact-col'><h3>⚙️ Assets</h3><ul>{''.join([f'<li>{x}</li>' for x in p.get('assets',[])])}</ul></div>
        <div class='pact-col'><h3>📋 Case Studies</h3><ul>{''.join([f'<li>{x}</li>' for x in p.get('case_studies',[])])}</ul></div>
        <div class='pact-col'><h3>💡 Thought Leadership</h3><ul>{''.join([f'<li>{x}</li>' for x in p.get('thought_leadership',[])])}</ul></div>
    </div>
    <div class='two'>
        <div>
            <h2>Objection Handling</h2>
            {''.join([f"<div class='obj'><strong>{o.get('q','')}</strong><p>{o.get('a','')}</p></div>" for o in p.get('objections',[])])}
        </div>
        <div>
            <h2>Stakeholder Hooks</h2>
            {''.join([f"<div class='hook'>{h}</div>" for h in p.get('stakeholder_hooks',[])])}
        </div>
    </div>
    <div class='two'>
        <div>
            <h2 style='text-align:center'>FAQs</h2>
            {''.join([f"<div class='faq'><strong>{f.get('q','')}</strong><p>{f.get('a','')}</p></div>" for f in p.get('faqs',[])])}
        </div>
        <div>
            <h2 style='text-align:center'>Key Market Insights</h2>
            {''.join([f"<div class='insight'>{i}</div>" for i in p.get('market_insights',[])])}
        </div>
    </div>
    </body></html>
    """

    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button("📂 Upload a Different Document", use_container_width=True):
            st.session_state.packet = None
            st.session_state.doc_name = ""
            st.session_state.doc_source = ""
            st.rerun()
    with col2:
        st.download_button(
            label="⬇️ Download Capability Packet",
            data=download_html,
            file_name=f"{p.get('title','capability-packet').replace(' ','-').lower()}.html",
            mime="text/html",
            use_container_width=True
        )